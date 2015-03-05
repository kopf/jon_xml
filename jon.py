import codecs
import os
import sys

from bs4 import BeautifulSoup
from docx import Document

FIELDS = ['title', 'metaTitle', 'fullTitle', 'shortTitle', 'metaDescription',
          'shortDescription', 'fullDescription', 'description']


def main(directory):
    for filename in [f for f in os.listdir(directory) if f.lower().endswith('.xml')]:
        doc = Document()
        with open(os.path.join(directory, filename), 'r') as f:
            xml = f.read()
        soup = BeautifulSoup(xml)
        for field in FIELDS:
            try:
                text = soup.find(field).text
            except (AttributeError, IndexError):
                continue
            if text.endswith(']]>'):
                text = text[:-3]
            doc.add_heading(field, 0)
            doc.add_paragraph(text)

        output_filename = os.path.join(directory, filename.lower().replace('.xml', '.docx'))
        doc.save(output_filename)

    print 'Done!'


if __name__ == '__main__':
    if not os.path.isdir(sys.argv[-1]):
        print 'Usage: '
        print 'jon.exe <path to directory containing XMLs>'
        sys.exit(-1)
    main(sys.argv[-1])
